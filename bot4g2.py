import asyncio
import logging
from aiogram import Bot, Dispatcher, types
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, BufferedInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.filters import Command
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.utils.keyboard import ReplyKeyboardBuilder
import psycopg2
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib import font_manager
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.graphics import renderPDF
from reportlab.graphics.shapes import Drawing, Rect, String
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.linecharts import LineChart
from reportlab.lib.colors import Color
from reportlab.lib.utils import ImageReader #для работы с буфером
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import io
from calendar import monthrange
import calendar
import seaborn as sns
from dotenv import load_dotenv
import os
from cryptography.fernet import Fernet

# Установка стиля seaborn для красивого оформления
sns.set_style("ticks")  # Белый фон с легкой сеткой
plt.style.use("seaborn-v0_8")  # Современный стиль seaborn

# Настройка шрифтов
plt.rcParams['font.family'] = 'DejaVu Sans'  # Или другой шрифт из available_fonts
plt.rcParams['font.size'] = 14  # Базовый размер шрифта
plt.rcParams['axes.titlesize'] = 18  # Размер заголовка
plt.rcParams['axes.labelsize'] = 14  # Размер меток осей
plt.rcParams['xtick.labelsize'] = 12  # Размер подписей на осях
plt.rcParams['ytick.labelsize'] = 12
plt.rcParams['legend.fontsize'] = 12
plt.rcParams['figure.figsize'] = [13, 7]  # Увеличенный размер графиков по умолчанию

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("bot4g2.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Настройка шрифтов для matplotlib
available_fonts = [f.name for f in font_manager.fontManager.ttflist]
preferred_fonts = ['DejaVu Sans', 'Arial', 'Times New Roman', 'Liberation Sans']
selected_font = next((font for font in preferred_fonts if font in available_fonts), None)

if selected_font:
    logger.info(f"Выбран шрифт для matplotlib: {selected_font}")
    matplotlib.rcParams['font.family'] = selected_font
else:
    try:
        font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
        font_manager.fontManager.addfont(font_path)
        matplotlib.rcParams['font.family'] = 'DejaVu Sans'
        logger.info("Шрифт DejaVuSans загружен вручную.")
    except Exception as e:
        logger.error(f"Не удалось загрузить шрифт DejaVuSans: {str(e)}")
        matplotlib.rcParams['font.family'] = 'sans-serif'

matplotlib.rcParams['font.size'] = 12

# Регистрация шрифта в reportlab
try:
    font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
    pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
    logger.info("Шрифт DejaVuSans зарегистрирован в reportlab.")
except Exception as e:
    logger.error(f"Не удалось зарегистрировать шрифт DejaVuSans в reportlab: {str(e)}")

# Загружаем переменные окружения из .env
load_dotenv()

# Получаем токен бота и ключ шифрования
API_TOKEN = os.getenv('API_TOKEN')
ENCRYPTION_KEY = os.getenv("ENCRYPTION_KEY")
cipher = Fernet(ENCRYPTION_KEY)

# Определяем функции шифрования
def encrypt_data(data: str) -> bytes:
    return cipher.encrypt(data.encode())

def decrypt_data(encrypted_data: bytes) -> str:
    return cipher.decrypt(encrypted_data).decode()

DB_CONFIG = {
    "dbname": os.getenv("DB_NAME"),
    "user": os.getenv("DB_USER"),
    "password": os.getenv("DB_PASSWORD"),
    "host": os.getenv("DB_HOST"),
    "port": os.getenv("DB_PORT"),
    "sslmode": "require"
}

# Инициализация бота
bot = Bot(token=API_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(bot=bot, storage=storage)

# Функция маскирует чувствительные данные 
def sanitize_log_data(data):
    if not data:
        return data
    sanitized_data = []
    for row in data:
        sanitized_row = []
        for item in row:
            if isinstance(item, str) and item not in ['Не указан', 'Без категории']:
                sanitized_row.append("***")
            else:
                sanitized_row.append(item)
        sanitized_data.append(tuple(sanitized_row))
    return sanitized_data

# Главное меню
main_menu = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text="Графики")],
        [KeyboardButton(text="Отчеты")],
        [KeyboardButton(text="Анализ товара")],
        [KeyboardButton(text="Анализ продаж")],
        [KeyboardButton(text="Помощь")],
    ],
    resize_keyboard=True
)

# SQL-запросы
SQL_QUERIES = {
    "sales_dynamics": 
        """SELECT 
            DATE_TRUNC('day', o."Date_order") AS "День",
            SUM(og."Sum_and_discont_og") AS "Общая выручка"
        FROM "Order" o
        LEFT JOIN "Order_goods" og ON o."OrderID" = og."OrderID"
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY DATE_TRUNC('day', o."Date_order")
        ORDER BY "День" ASC;""",
    "category_sales": 
        """SELECT 
            COALESCE(cg."Category", 'Без категории') AS "Категория товара",
            SUM(og."Sum_and_discont_og") AS "Выручка по категории"
        FROM public."Order_goods" og
        LEFT JOIN public."Goods" g ON og."GoodID" = g."GoodID"
        LEFT JOIN public."Category_goods" cg ON g."Category_goodsID" = cg."Category_goodsID"
        LEFT JOIN public."Order" o ON og."OrderID" = o."OrderID"
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY cg."Category"
        ORDER BY "Выручка по категории" DESC;""",
    "city_revenue":
        """SELECT 
            COALESCE(
                CASE 
                    WHEN o."Buying_method" = 'Онлайн' AND o."DeliveriID" != 0 AND d."AdressID" != 0 THEN a."City"
                    ELSE s."City"
                END, 
                'Не указан'
            ) AS "Город",
            SUM(og."Sum_and_discont_og") AS "Выручка"
        FROM public."Order" o
        LEFT JOIN public."Realization" r ON o."RealizationID" = r."RealizationID"
        LEFT JOIN public."Store" s ON r."StoreID" = s."StoreID"
        LEFT JOIN public."Delivery" d ON o."DeliveriID" = d."DeliveryID"
        LEFT JOIN public."Address" a ON d."AdressID" = a."AddressID"
        LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY 
            CASE 
                WHEN o."Buying_method" = 'Онлайн' AND o."DeliveriID" != 0 AND d."AdressID" != 0 THEN a."City"
                ELSE s."City"
            END
        ORDER BY "Выручка" DESC
        LIMIT 19;""",
    "payment_methods": 
        """SELECT 
            p."Method_payment",
            COUNT(DISTINCT o."OrderID") AS "Количество заказов"
        FROM "Order" o
        JOIN "Payment" p ON o."PaymentID" = p."PaymentID"
        JOIN "Order_goods" og ON o."OrderID" = og."OrderID"
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY p."Method_payment"
        ORDER BY "Количество заказов" DESC;""",
    "gender_stats": 
        """SELECT 
            CASE 
                WHEN c."Gender" IS NULL THEN 'Не указан'
                ELSE c."Gender"
            END AS "Пол",
            COUNT(c."CustomerID") AS "Количество покупателей"
        FROM public."Customer" c
        JOIN public."Order" o ON c."CustomerID" = o."CustomerID"
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY c."Gender"
        ORDER BY "Количество покупателей" DESC;""",
    "top_goods": 
        """SELECT 
            g."Goods" AS "Название товара",
            COALESCE(SUM(og."Quantity_goods"), 0) AS "Количество проданных единиц",
            COALESCE(SUM(og."Sum_and_discont_og"), 0) AS "Общая выручка"
        FROM public."Goods" g
        JOIN public."Order_goods" og ON g."GoodID" = og."GoodID"
        JOIN public."Order" o ON og."OrderID" = o."OrderID" AND o."Order status" = 'Завершен'
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY g."Goods"
        ORDER BY "Количество проданных единиц" DESC, "Общая выручка" DESC
        LIMIT 10;""",
    "order_dynamics": 
        """SELECT 
            DATE_TRUNC('day', o."Date_order") AS "День",
            COUNT(o."OrderID") AS "Количество заказов"
        FROM "Order" o
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY DATE_TRUNC('day', o."Date_order")
        ORDER BY "День" ASC;""",
    "top_brend":
    """SELECT 
            g."Brend" AS "Бренд",
            COALESCE(SUM(og."Quantity_goods"), 0) AS "Количество проданных единиц",
            COALESCE(SUM(og."Sum_and_discont_og"), 0) AS "Общая выручка"
        FROM public."Goods" g
        JOIN public."Order_goods" og ON g."GoodID" = og."GoodID"
        JOIN public."Order" o ON og."OrderID" = o."OrderID" AND o."Order status" = 'Завершен'
        WHERE o."Date_order" BETWEEN %s AND %s
        GROUP BY g."Brend"
        ORDER BY "Количество проданных единиц" DESC, "Общая выручка" DESC
        LIMIT 15;
    """
}

# Функция подключения к БД
def get_db_connection():
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        logger.info("Успешное подключение к базе данных.")
        return conn
    except Exception as e:
        logger.error(f"Ошибка подключения к базе данных: {str(e)}")
        raise

# Функция создания графика
def create_graph(query_name, start_date, end_date):
    try:
        # Подключение к базе данных
        conn = get_db_connection()
        cur = conn.cursor()
        # Выполняем SQL-запрос
        cur.execute(SQL_QUERIES[query_name], (start_date, end_date))
        data = cur.fetchall()
        
        if not data:
            logger.warning(f"Нет данных для графика '{query_name}' за период {start_date} - {end_date}")
            cur.close()
            conn.close()
            return None, f"Нет данных для графика '{query_name}' за период {start_date} - {end_date}."
        
        logger.info(f"Данные для графика '{query_name}': {data}")
        
        # Создаём фигуру и оси
        fig, ax = plt.subplots(figsize=(14, 8), dpi=300)
        colors = sns.color_palette("rocket", n_colors=19)
        
      # Динамика выручки
        if query_name == "sales_dynamics":
            x = [row[0] for row in data]
            y = [row[1] / 1000 for row in data]  # Выручка в тысячах
            # Используем цвет #e95150 для выручки, уменьшаем толщину линии до 2
            ax.plot(x, y, color="#e95150", linewidth=2, marker='o', markersize=8, label='Выручка')
            ax.fill_between(x, y, alpha=0.15, color="#e95150")
            # Добавляем серую клеточную сетку
            ax.grid(True, color='gray', linestyle='-', linewidth=0.5, alpha=0.7)
            ax.grid(True, color='gray', linestyle='-', linewidth=0.3, alpha=0.5)
            ax.set_title(f"Динамика выручки ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            ax.set_xlabel("Дата", labelpad=15, fontsize=14)
            ax.set_ylabel("Выручка, тыс. ₽", labelpad=15, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=12)
            ax.legend(loc='upper left', frameon=True, shadow=True, fontsize=14)
        
      # Категории товаров
        elif query_name == "category_sales":
            x = [row[0] for row in data]
            y = [float(row[1]) / 1000 for row in data]  # Выручка в тысячах
            # Уменьшаем width до 0.3 для увеличения пространства между столбцами
            bars = ax.bar(x, y, color=colors, edgecolor='black', linewidth=0.8, alpha=0.85, width=0.3)
            for bar in bars:
                bar.set_zorder(2)
                bar.set_edgecolor("grey")
            ax.set_title(f"Выручка по категориям ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=30, fontsize=20)
            ax.set_xlabel("Категория", labelpad=15, fontsize=14)
            ax.set_ylabel("Выручка, тыс. ₽", labelpad=15, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=12)
            max_y = max(y)
            for bar in bars:
                yval = bar.get_height()
                offset = max_y * 0.05
                # Отображаем значение с одной цифрой после запятой, под углом 25 градусов
                ax.text(bar.get_x() + bar.get_width()/2, yval + offset, f"{yval:.1f}", 
                        ha='center', va='bottom', fontsize=12, fontweight='bold', rotation=25)
        
        # Выручка по городам
        elif query_name == "city_revenue":
            x = [row[0] for row in data]
            y = [float(row[1]) / 1000 for row in data]  # Выручка в тысячах
            # Уменьшаем width до 0.3 для увеличения пространства между столбцами
            bars = ax.bar(x, y, color=colors, edgecolor='black', linewidth=0.8, alpha=0.85, width=0.3)
            for bar in bars:
                bar.set_zorder(2)
                bar.set_edgecolor("grey")
            # Увеличиваем pad до 30, чтобы поднять заголовок выше
            ax.set_title(f"Выручка по городам ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=35, fontsize=20)
            ax.set_xlabel("Город", labelpad=15, fontsize=14)
            ax.set_ylabel("Выручка, тыс. ₽", labelpad=15, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=12)
            max_y = max(y)
            for bar in bars:
                yval = bar.get_height()
                offset = max_y * 0.05
                # Отображаем значение с одной цифрой после запятой, под углом 25 градусов
                ax.text(bar.get_x() + bar.get_width()/2, yval + offset, f"{yval:.1f}", 
                        ha='center', va='bottom', fontsize=12, fontweight='bold', rotation=25)
        
        # Методы оплаты
        elif query_name == "payment_methods":
            labels = [row[0] for row in data]
            sizes = [row[1] for row in data]
            explode = [0.05] * len(sizes)
            pie_colors = ["#953269", "#e95150", "#f5936e"]
            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=pie_colors, 
                   shadow=False, explode=explode, textprops={'fontsize': 16, 'color': 'white', 'fontweight': 'bold'})
            ax.set_title(f"Методы оплаты ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            ax.legend(labels=['Наличные', 'Карта'], loc='upper right', fontsize=16)
        
        # Круговая диаграмма для распределения по полу
        elif query_name == "gender_stats":
            # Извлекаем данные: labels — категории (пол), sizes — значения (процент или количество)
            labels = [row[0] for row in data]
            sizes = [row[1] for row in data]
            # explode задаёт "выдвижение" сегментов (0.05 — небольшое выдвижение для всех)
            explode = [0.05] * len(sizes)
            pie_colors = ["#953269", "#e95150", "#f5936e"]
            # Строим круговую диаграмму
            ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=pie_colors, 
                shadow=False, explode=explode, textprops={'fontsize': 16, 'color': 'white', 'fontweight': 'bold'})
            # Заголовок графика с форматированием дат
            ax.set_title(f"Распределение по полу ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            # Добавляем легенду
            ax.legend(labels=['Ж', 'М', 'Не указан'], loc='upper right', fontsize=16)
        
        # Динамика заказов
        elif query_name == "order_dynamics":
            x = [row[0] for row in data]
            y = [row[1] for row in data]  # Количество заказов
            # Используем цвет #e95150 для заказов, уменьшаем толщину линии до 2
            ax.plot(x, y, color="#e95150", linewidth=2, marker='o', markersize=8, label='Количество заказов')
            ax.fill_between(x, y, alpha=0.15, color="#e95150")
            # Добавляем серую клеточную сетку (major grid)
            ax.grid(True, which='major', color='gray', linestyle='-', linewidth=0.5, alpha=0.7)
            ax.set_title(f"Динамика заказов ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            ax.set_xlabel("Дата", labelpad=15, fontsize=14)
            ax.set_ylabel("Количество заказов", labelpad=15, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=12)
            ax.legend(loc='upper left', frameon=True, shadow=True, fontsize=14)
        
        # Топ-10 товаров
        elif query_name == "top_goods":
            # Увеличиваем высоту графика (ширина 10 дюймов, высота 8 дюймов)
            fig.set_size_inches(12, 13)
            x = [row[0] for row in data]
            y = [float(row[1]) for row in data]  # Количество проданных единиц
            bars = ax.bar(x, y, color=colors, edgecolor='black', linewidth=0.8, alpha=0.85, width=0.4)
            for bar in bars:
                bar.set_zorder(2)
                bar.set_edgecolor("grey")
            ax.set_title(f"Топ-10 товаров ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            ax.set_xlabel("Товар", labelpad=15, fontsize=14)
            ax.set_ylabel("Количество проданных единиц", labelpad=20, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=10)
            max_y = max(y)
            for bar in bars:
                yval = bar.get_height()
                offset = max_y * 0.05
                ax.text(bar.get_x() + bar.get_width()/2, yval + offset, f"{int(yval)}", 
                        ha='center', va='bottom', fontsize=12, fontweight='bold')
        
        # Топ-15 брендов
        elif query_name == "top_brend":
            x = [row[0] for row in data]
            y = [float(row[1]) for row in data]  # Количество проданных единиц
            bars = ax.bar(x, y, color=colors, edgecolor='black', linewidth=0.8, alpha=0.85, width=0.4)
            for bar in bars:
                bar.set_zorder(2)
                bar.set_edgecolor("grey")
            ax.set_title(f"Топ-15 брендов ({start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')})", pad=20, fontsize=20)
            ax.set_xlabel("Бренд", labelpad=15, fontsize=14)
            ax.set_ylabel("Количество проданных единиц", labelpad=15, fontsize=14)
            plt.xticks(rotation=45, ha='right', fontsize=10)
            max_y = max(y)
            for bar in bars:
                yval = bar.get_height()
                offset = max_y * 0.05
                ax.text(bar.get_x() + bar.get_width()/2, yval + offset, f"{int(yval)}", 
                        ha='center', va='bottom', fontsize=12, fontweight='bold')
        
        # Общие настройки
        ax.set_facecolor('#f8f8f8')
        fig.patch.set_facecolor('#ffffff')
        ax.grid(True, linestyle='--', alpha=0.7, zorder=0)
        ax.tick_params(axis='both', which='major', labelsize=12)
        plt.tight_layout()
        
        # Сохраняем график
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        cur.close()
        conn.close()
        return buffer, None
    except Exception as e:
        logger.error(f"Ошибка при создании графика '{query_name}': {str(e)}")
        return None, f"Ошибка при создании графика: {str(e)}"

# Функция создания PDF для дашборда
def create_pdf(query_name, graph_buffer, data, column_names, start_date, end_date):
    try:
        # Если это дашборд, просто возвращаем буфер, так как он уже в формате PDF
        if query_name == "Дашборд":
            return graph_buffer

        # Для других типов отчётов (оставляем старую логику)
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=landscape(letter))
        styles = getSampleStyleSheet()
        
        styles['Title'].fontName = 'DejaVuSans'
        styles['Title'].fontSize = 16
        styles['Title'].leading = 20
        
        story = []
        
        story.append(Paragraph(f"Отчет: {query_name} за {start_date} - {end_date}", styles['Title']))
        story.append(Spacer(1, 12))
        
        img = Image(graph_buffer, width=700, height=400)
        story.append(img)
        
        story.append(PageBreak())
        
        story.append(Paragraph(f"Данные отчета за {start_date} - {end_date}", styles['Title']))
        story.append(Spacer(1, 12))
        
        table_data = [column_names]
        for row in data:
            table_data.append([str(cell) for cell in row])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'DejaVuSans'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'DejaVuSans'),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(table)
        
        doc.build(story)
        pdf_buffer.seek(0)
        return pdf_buffer
    except Exception as e:
        logger.error(f"Ошибка при создании PDF для '{query_name}': {str(e)}")
        return None

def create_dashboard(start_date, end_date):
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Создаём PDF с размером страницы
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=(2318, 1980))
        width, height = 2318, 1980

        # Устанавливаем фон страницы (темнее: #f5f5f5)
        c.setFillColor(Color(0.956, 0.956, 0.956))  # f5f5f5
        c.rect(0, 0, width, height, fill=1, stroke=0)

        # --- Заголовки и метрики ---
        c.setFont("DejaVuSans", 24)
        c.setFillColor(Color(0, 0, 0))  # Чёрный текст

        # Параметры закруглённого прямоугольника под заголовком "Дашборд"
        c.setFillColor(Color(1, 1, 1))  # Белый фон для прямоугольника
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))  # Серая обводка
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(0, height - 25 - 12 - 30, 510, 50, radius=10, stroke=1, fill=1)
        # Текст заголовка
        c.setFillColor(Color(0, 0, 0))  # Чёрный текст
        c.drawString(15, height - 35 - 12, f"Дашборд ({start_date.strftime('%Y-%m-%d')} - {end_date.strftime('%Y-%m-%d')})")

        # Получаем данные для метрик
        table_data = get_dashboard_table_data(start_date, end_date)
        logger.info(f"Данные для заголовков: {table_data}")
        total_revenue = table_data[0][1] if table_data else 0
        order_count = table_data[1][1] if table_data else 0
        avg_check = table_data[2][1] if table_data else 0

        # Параметры закруглённого прямоугольника под метриками (единый для всех трёх)
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(790, height - 25 - 12 - 30, 1420, 50, radius=10, stroke=1, fill=1)
        # Метрики
        c.setFillColor(Color(0, 0, 0))
        c.drawString(800, height - 35 - 12, f"Общая выручка: {total_revenue} ₽")
        c.drawString(1300, height - 35 - 12, f"Кол-во заказов: {order_count}")
        c.drawString(1750, height - 35 - 12, f"Средний чек: {avg_check} ₽")

        # --- Панельный график (Динамика выручки) 1 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(53, height - 165 - 358 - 20, 2165 + 20, 358 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(50, height - 119 - 12, "Динамика выручки")
        cur.execute(SQL_QUERIES["sales_dynamics"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для sales_dynamics: {data}")
        if data:
            days = [row[0] for row in data]
            sales = [float(row[1]) if row[1] is not None else 0 for row in data]
            # Параметры графика: размер графика в matplotlib (ширина, высота в дюймах)
            fig, ax = plt.subplots(figsize=(15.51, 3.58), dpi=100)
            # Настройка стиля графика
            ax.plot(days, sales, color='blue', linewidth=2, marker='o', markersize=6)
            ax.set_xlabel("Дата продажи", labelpad=10, fontfamily='DejaVu Sans', fontsize=10)
            ax.set_ylabel("Выручка, ₽", labelpad=10, fontfamily='DejaVu Sans', fontsize=10)
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%d.%m'))
            plt.xticks(rotation=45, ha='right', fontfamily='DejaVu Sans', fontsize=8)
            ax.grid(True, linestyle='--', alpha=0.7)
            plt.tight_layout()
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=100, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            img_reader = ImageReader(buffer)
            # Параметры вставки графика: x, y, ширина, высота
            c.drawImage(img_reader, 77, height - 165 - 358, width=2135, height=366)
        else:
            c.drawString(169 + 1551/2, height - 165 - 358 + 358/2, "Нет данных", fontName="DejaVuSans", fontSize=12)

        # --- Панельный график (Динамика заказов) 2 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(53, height - 815 - 347 - 10, 710 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(53, height - 596 - 12, "Динамика заказов")
        cur.execute(SQL_QUERIES["order_dynamics"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для order_dynamics: {data}")
        if data:
            days = [row[0] for row in data]
            orders = [int(row[1]) if row[1] is not None else 0 for row in data]
            # Параметры графика: размер графика в matplotlib (ширина, высота в дюймах)
            fig, ax = plt.subplots(figsize=(4.93, 2.97), dpi=100)
            # Настройка стиля графика
            ax.plot(days, orders, color='blue', linewidth=2, marker='o', markersize=6)
            ax.set_xlabel("Дата заказа", labelpad=10, fontfamily='DejaVu Sans', fontsize=10)
            ax.set_ylabel("Количество", labelpad=10, fontfamily='DejaVu Sans', fontsize=10)
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%d.%m'))
            plt.xticks(rotation=45, ha='right', fontfamily='DejaVu Sans', fontsize=8)
            ax.grid(True, linestyle='--', alpha=0.7)
            plt.tight_layout()
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=100, bbox_inches='tight')
            buffer.seek(0)
            plt.close()
            img_reader = ImageReader(buffer)
            # Параметры вставки графика: x, y, ширина, высота
            c.drawImage(img_reader, 60, height - 650 - 517, width=645, height=517)
        else:
            c.drawString(500 + 645/2, height - 660 - 347 + 347/2, "Нет данных", fontName="DejaVuSans", fontSize=12)

        # --- Столбчатая диаграмма (Выручка по городам) 3 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(844, height - 815 - 347 - 10, 710 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(850, height - 596 - 12, "Выручка по городам")
        # Параметры графика: размер области для графика (ширина, высота в пунктах)
        drawing = Drawing(645, 347)
        cur.execute(SQL_QUERIES["city_revenue"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для city_revenue: {data}")
        if data:
            cities = [row[0] for row in data]
            revenue = [float(row[1]) if row[1] is not None else 0 for row in data]
            bc = VerticalBarChart()
            # Настройка стиля графика
            bc.x = 0
            bc.y = 0
            bc.width = 645
            bc.height = 347
            bc.data = [revenue]
            bc.bars[0].fillColor = Color(0.957, 0.957, 0.957)
            bc.valueAxis.valueMin = 0
            bc.valueAxis.valueMax = max(revenue) * 1.1 if revenue else 100
            bc.categoryAxis.categoryNames = cities
            bc.categoryAxis.labels.angle = 45
            bc.categoryAxis.labels.boxAnchor = 'ne'
            drawing.add(bc)
        else:
            drawing.add(String(645/2, 347/2, "Нет данных", fontName="DejaVuSans", fontSize=12, textAnchor="middle"))
        # Параметры вставки графика: x, y
        renderPDF.draw(drawing, c, 920, height - 700 - 347)

        # --- Столбчатая диаграмма (Топ категорий) 6 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(830, height - 1480 - 347 - 10, 710 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(855, height - 1255 - 12, "Топ категорий")
        # Параметры графика: размер области для графика (ширина, высота в пунктах)
        drawing = Drawing(645, 347)
        cur.execute(SQL_QUERIES["category_sales"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для category_sales: {data}")
        if data:
            categories = [row[0] for row in data]
            revenue = [float(row[1]) if row[1] is not None else 0 for row in data]
            bc = VerticalBarChart()
            # Настройка стиля графика
            bc.x = 0
            bc.y = 0
            bc.width = 645
            bc.height = 347
            bc.data = [revenue]
            bc.bars[0].fillColor = Color(0.957, 0.957, 0.957)
            bc.valueAxis.valueMin = 0
            bc.valueAxis.valueMax = max(revenue) * 1.1 if revenue else 100
            bc.categoryAxis.categoryNames = categories
            bc.categoryAxis.labels.angle = 45
            bc.categoryAxis.labels.boxAnchor = 'ne'
            drawing.add(bc)
        else:
            drawing.add(String(645/2, 347/2, "Нет данных", fontName="DejaVuSans", fontSize=12, textAnchor="middle"))
        # Параметры вставки графика: x, y
        renderPDF.draw(drawing, c, 920, height - 1330 - 347)

        # --- Столбчатая диаграмма (Топ товаров) 5 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(53, height - 1480 - 347 - 20, 710 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(53, height - 1250 - 12, "Топ товаров")
        # Параметры графика: размер области для графика (ширина, высота в пунктах)
        drawing = Drawing(645, 347)
        cur.execute(SQL_QUERIES["top_goods"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для top_goods: {data}")
        if data:
            goods = [row[0] for row in data]
            quantities = [int(row[1]) if row[1] is not None else 0 for row in data]
            bc = VerticalBarChart()
            # Настройка стиля графика
            bc.x = 0
            bc.y = 0
            bc.width = 645
            bc.height = 347
            bc.data = [quantities]
            bc.bars[0].fillColor = Color(0.957, 0.957, 0.957)
            bc.valueAxis.valueMin = 0
            bc.valueAxis.valueMax = max(quantities) * 1.1 if quantities else 100
            bc.categoryAxis.categoryNames = goods
            bc.categoryAxis.labels.angle = 45
            bc.categoryAxis.labels.boxAnchor = 'ne'
            drawing.add(bc)
        else:
            drawing.add(String(53, 347/2, "Нет данных", fontName="DejaVuSans", fontSize=12, textAnchor="middle"))
        # Параметры вставки графика: x, y
        renderPDF.draw(drawing, c, 110, height - 1330 - 347)

        # --- Круговой график (Методы оплаты) 4 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(1710, height - 860 - 300 - 10, 510 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(1710, height - 596 - 12, "Методы оплаты")
        # Параметры графика: размер области для графика (ширина, высота в пунктах)
        drawing = Drawing(300, 300)
        cur.execute(SQL_QUERIES["payment_methods"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для payment_methods: {data}")
        if data:
            labels = [row[0] for row in data]
            sizes = [int(row[1]) if row[1] is not None else 0 for row in data]
            pie = Pie()
            # Настройка стиля графика
            pie.x = 0
            pie.y = 0
            pie.width = 400
            pie.height = 400
            pie.data = sizes
            pie.labels = labels
            pie.slices.strokeColor = None
            pie.slices[0].fillColor = Color(0.957, 0.957, 0.957)
            pie.slices[1].fillColor = Color(0.8, 0.8, 0.8)
            pie.slices[2].fillColor = Color(0.6, 0.6, 0.6)
            drawing.add(pie)
        else:
            drawing.add(String(300/2, 300/2, "Нет данных", fontName="DejaVuSans", fontSize=12, textAnchor="middle"))
        # Параметры вставки графика: x, y
        renderPDF.draw(drawing, c, 1780, height - 760 - 300)

        # --- Круговой график (Распределение по гендеру) 7 график ---
        # Параметры закруглённого прямоугольника под графиком
        c.setFillColor(Color(1, 1, 1))
        #c.setStrokeColor(Color(0.5, 0.5, 0.5))
        # Параметры прямоугольника: x, y, ширина, высота, радиус скругления
        c.roundRect(1710, height - 1530 - 300 - 10, 510 + 20, 510 + 40, radius=15, stroke=1, fill=1)
        # Заголовок
        c.setFillColor(Color(0, 0, 0))
        c.drawString(1710, height - 1260 - 12, "Распределение по гендеру")
        # Параметры графика: размер области для графика (ширина, высота в пунктах)
        drawing = Drawing(300, 300)
        cur.execute(SQL_QUERIES["gender_stats"], (start_date, end_date))
        data = cur.fetchall()
        logger.info(f"Данные для gender_stats: {data}")
        if data:
            labels = [row[0] for row in data]
            sizes = [int(row[1]) if row[1] is not None else 0 for row in data]
            pie = Pie()
            # Настройка стиля графика
            pie.x = 0
            pie.y = 0
            pie.width = 400
            pie.height = 400
            pie.data = sizes
            pie.labels = labels
            pie.slices.strokeColor = None
            pie.slices[0].fillColor = Color(0.957, 0.957, 0.957)
            pie.slices[1].fillColor = Color(0.8, 0.8, 0.8)
            pie.slices[2].fillColor = Color(0.6, 0.6, 0.6)
            drawing.add(pie)
        else:
            drawing.add(String(300/2, 300/2, "Нет данных", fontName="DejaVuSans", fontSize=12, textAnchor="middle"))
        # Параметры вставки графика: x, y
        renderPDF.draw(drawing, c, 1780, height - 1440 - 300)

        # Сохраняем PDF
        c.showPage()
        c.save()
        pdf_buffer.seek(0)

        cur.close()
        conn.close()
        return pdf_buffer

    except Exception as e:
        logger.error(f"Ошибка при создании дашборда: {str(e)}")
        return None
    
# Функция для получения данных таблицы для дашборда
def get_dashboard_table_data(start_date, end_date):
    try:
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT 
                'Общая выручка' AS "Показатель", 
                COALESCE(SUM(og."Sum_and_discont_og"), 0) AS "Значение"
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            UNION ALL
            SELECT 
                'Количество заказов', 
                COUNT(DISTINCT o."OrderID")
            FROM public."Order" o
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            UNION ALL
            SELECT 
                'Средний чек', 
                ROUND(COALESCE(SUM(og."Sum_and_discont_og"), 0) / NULLIF(COUNT(DISTINCT o."OrderID"), 0), 2)
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (start_date, end_date, start_date, end_date, start_date, end_date))
        data = cur.fetchall()
        cur.close()
        conn.close()
        return data
    except Exception as e:
        logger.error(f"Ошибка при получении данных для дашборда: {str(e)}")
        return []

# Функции для отчетов
def get_weekly_report_data(start_date, end_date):
    logger.info(f"get_weekly_report_data: start_date={start_date}, type={type(start_date)}, end_date={end_date}, type={type(end_date)}")
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Преобразуем даты в строки
        start_date_str = start_date.strftime('%Y-%m-%d')
        end_date_str = end_date.strftime('%Y-%m-%d')

        # Total revenue
        cur.execute("""
            SELECT COALESCE(SUM(og."Sum_and_discont_og"), 0) AS total_revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (start_date_str, end_date_str))
        total_revenue = cur.fetchone()[0] or 0

        # Sales count
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS sales_count
            FROM public."Order" o
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (start_date_str, end_date_str))
        sales_count = cur.fetchone()[0] or 0

        # Average check
        avg_check = total_revenue / sales_count if sales_count > 0 else 0

        # Sales dynamics
        prev_start_date = start_date - timedelta(days=7)
        prev_end_date = end_date - timedelta(days=7)
        prev_start_date_str = prev_start_date.strftime('%Y-%m-%d')
        prev_end_date_str = prev_end_date.strftime('%Y-%m-%d')
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS prev_sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) AS prev_revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (prev_start_date_str, prev_end_date_str))
        prev_sales_count, prev_revenue = cur.fetchone() or (0, 0)
        sales_dynamics = ((sales_count - prev_sales_count) / prev_sales_count * 100) if prev_sales_count > 0 else 0

        # New customers
        cur.execute("""
            SELECT COUNT(*) AS new_customers
            FROM public."Customer" c
            WHERE c."Registration_date" BETWEEN %s AND %s;
        """, (start_date_str, end_date_str))
        new_customers = cur.fetchone()[0] or 0

        # Top products
        cur.execute("""
            SELECT g."Goods", SUM(og."Quantity_goods") AS quantity_sold, SUM(og."Sum_and_discont_og") AS revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            LEFT JOIN public."Goods" g ON og."GoodID" = g."GoodID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY g."Goods"
            ORDER BY quantity_sold DESC
            LIMIT 5;
        """, (start_date_str, end_date_str))
        top_products = cur.fetchall() or []

        # Channels (Online vs Offline)
        cur.execute("""
            SELECT o."Buying_method" AS channel, 
                   COUNT(DISTINCT o."OrderID") AS sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) AS revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY o."Buying_method";
        """, (start_date_str, end_date_str))
        channel_data = cur.fetchall() or []
        channels = [(row[0], row[1], row[2]) for row in channel_data]
        channels.append(("Итог", sales_count, total_revenue))

        # Daily data
        cur.execute("""
            SELECT o."Date_order"::date AS sale_date,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) AS revenue,
                   COUNT(DISTINCT o."OrderID") AS sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) / NULLIF(COUNT(DISTINCT o."OrderID"), 0) AS avg_check
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY o."Date_order"::date
            ORDER BY sale_date;
        """, (start_date_str, end_date_str))
        daily_data_raw = cur.fetchall() or []

        daily_data = []
        prev_week_start = start_date - timedelta(days=7)
        prev_week_end = end_date - timedelta(days=7)
        prev_week_start_str = prev_week_start.strftime('%Y-%m-%d')
        prev_week_end_str = prev_week_end.strftime('%Y-%m-%d')
        cur.execute("""
            SELECT o."Date_order"::date AS sale_date,
                   COUNT(DISTINCT o."OrderID") AS prev_sales_count
            FROM public."Order" o
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY o."Date_order"::date
            ORDER BY sale_date;
        """, (prev_week_start_str, prev_week_end_str))
        prev_daily_sales = {row[0]: row[1] for row in cur.fetchall()}

        current_date = start_date
        while current_date <= end_date:
            found = False
            for day_data in daily_data_raw:
                if day_data[0] == current_date:
                    revenue, sales, avg_check = day_data[1], day_data[2], day_data[3] or 0
                    prev_sales = prev_daily_sales.get(current_date - timedelta(days=7), 0)
                    change = ((sales - prev_sales) / prev_sales * 100) if prev_sales > 0 else 0
                    daily_data.append((current_date, revenue, sales, avg_check, change))
                    found = True
                    break
            if not found:
                daily_data.append((current_date, 0, 0, 0, 0))
            current_date += timedelta(days=1)

        # Delivery data
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS shipped_orders
            FROM public."Order" o
            LEFT JOIN public."Delivery" d ON o."DeliveriID" = d."DeliveryID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен' AND d."DeliveryID" IS NOT NULL;
        """, (start_date_str, end_date_str))
        shipped_orders = cur.fetchone()[0] or 0

        cur.execute("""
            SELECT COALESCE(STRING_AGG(DISTINCT a."City", ', '), 'Не указан') AS main_regions
            FROM public."Order" o
            LEFT JOIN public."Delivery" d ON o."DeliveriID" = d."DeliveryID"
            LEFT JOIN public."Address" a ON d."AdressID" = a."AddressID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен' AND a."City" IS NOT NULL
            LIMIT 3;
        """, (start_date_str, end_date_str))
        main_regions = cur.fetchone()[0] or "Москва, Санкт-Петербург, Казань"

        avg_delivery_time = "2 дня"  # Placeholder

        cur.close()
        conn.close()

        return {
            "total_revenue": total_revenue,
            "sales_count": sales_count,
            "avg_check": avg_check,
            "sales_dynamics": sales_dynamics,
            "new_customers": new_customers,
            "top_products": [(row[0], row[1]) for row in top_products],
            "top_product_revenue": top_products[0][2] if top_products else 0,
            "channels": channels,
            "daily_data": daily_data,
            "shipped_orders": shipped_orders,
            "avg_delivery_time": avg_delivery_time,
            "main_regions": main_regions
        }
    except Exception as e:
        logger.error(f"Ошибка при получении данных для еженедельного отчета: {str(e)}")
        return {
            "total_revenue": 0,
            "sales_count": 0,
            "avg_check": 0,
            "sales_dynamics": 0,
            "new_customers": 0,
            "top_products": [],
            "top_product_revenue": 0,
            "channels": [],
            "daily_data": [],
            "shipped_orders": 0,
            "avg_delivery_time": "0",
            "main_regions": "Не указан"
        }

def get_monthly_report_data(start_date, end_date):
    logger.info(f"get_monthly_report_data: start_date={start_date}, type={type(start_date)}, end_date={end_date}, type={type(end_date)}")
    try:
        conn = get_db_connection()
        cur = conn.cursor()

        # Преобразуем даты в строки
        start_date_str = start_date.strftime('%Y-%m-%d')
        end_date_str = end_date.strftime('%Y-%m-%d')

        # Total revenue
        cur.execute("""
            SELECT COALESCE(SUM(og."Sum_and_discont_og"), 0) AS total_revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (start_date_str, end_date_str))
        total_revenue = cur.fetchone()[0] or 0

        # Sales count
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS sales_count
            FROM public."Order" o
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (start_date_str, end_date_str))
        sales_count = cur.fetchone()[0] or 0

        # Average check
        avg_check = total_revenue / sales_count if sales_count > 0 else 0

        # Sales dynamics
        prev_start_date = start_date - timedelta(days=30)
        prev_end_date = end_date - timedelta(days=30)
        prev_start_date_str = prev_start_date.strftime('%Y-%m-%d')
        prev_end_date_str = prev_end_date.strftime('%Y-%m-%d')
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS prev_sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) AS prev_revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (prev_start_date_str, prev_end_date_str))
        prev_sales_count, prev_revenue = cur.fetchone() or (0, 0)
        sales_dynamics = ((sales_count - prev_sales_count) / prev_sales_count * 100) if prev_sales_count > 0 else 0

        # New customers
        cur.execute("""
            SELECT COUNT(*) AS new_customers
            FROM public."Customer" c
            WHERE c."Registration_date" BETWEEN %s AND %s;
        """, (start_date_str, end_date_str))
        new_customers = cur.fetchone()[0] or 0

        # Top products
        cur.execute("""
            SELECT g."Goods", SUM(og."Quantity_goods") AS quantity_sold, SUM(og."Sum_and_discont_og") AS revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            LEFT JOIN public."Goods" g ON og."GoodID" = g."GoodID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY g."Goods"
            ORDER BY quantity_sold DESC
            LIMIT 5;
        """, (start_date_str, end_date_str))
        top_products = cur.fetchall() or []

        # Channels (Online vs Offline)
        cur.execute("""
            SELECT o."Buying_method" AS channel, 
                   COUNT(DISTINCT o."OrderID") AS sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) AS revenue
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен'
            GROUP BY o."Buying_method";
        """, (start_date_str, end_date_str))
        channel_data = cur.fetchall() or []
        channels = [(row[0], row[1], row[2]) for row in channel_data]
        channels.append(("Итог", sales_count, total_revenue))

        # Monthly data (for dynamics comparison with previous month)
        cur.execute("""
            SELECT COALESCE(SUM(og."Sum_and_discont_og"), 0) AS revenue,
                   COUNT(DISTINCT o."OrderID") AS sales_count,
                   COALESCE(SUM(og."Sum_and_discont_og"), 0) / NULLIF(COUNT(DISTINCT o."OrderID"), 0) AS avg_check
            FROM public."Order" o
            LEFT JOIN public."Order_goods" og ON o."OrderID" = og."OrderID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен';
        """, (prev_start_date_str, prev_end_date_str))
        prev_month_data = cur.fetchone() or (0, 0, 0)
        monthly_data = [
            (prev_start_date, prev_month_data[0], prev_month_data[1], prev_month_data[2] or 0, 0),
            (start_date, total_revenue, sales_count, avg_check, sales_dynamics)
        ]

        # Delivery data
        cur.execute("""
            SELECT COUNT(DISTINCT o."OrderID") AS shipped_orders
            FROM public."Order" o
            LEFT JOIN public."Delivery" d ON o."DeliveriID" = d."DeliveryID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен' AND d."DeliveryID" IS NOT NULL;
        """, (start_date_str, end_date_str))
        shipped_orders = cur.fetchone()[0] or 0

        cur.execute("""
            SELECT COALESCE(STRING_AGG(DISTINCT a."City", ', '), 'Не указан') AS main_regions
            FROM public."Order" o
            LEFT JOIN public."Delivery" d ON o."DeliveriID" = d."DeliveryID"
            LEFT JOIN public."Address" a ON d."AdressID" = a."AddressID"
            WHERE o."Date_order" BETWEEN %s AND %s AND o."Order status" = 'Завершен' AND a."City" IS NOT NULL
            LIMIT 3;
        """, (start_date_str, end_date_str))
        main_regions = cur.fetchone()[0] or "Москва, Санкт-Петербург, Казань"

        avg_delivery_time = "2 дня"  # Placeholder

        cur.close()
        conn.close()

        return {
            "total_revenue": total_revenue,
            "sales_count": sales_count,
            "avg_check": avg_check,
            "sales_dynamics": sales_dynamics,
            "new_customers": new_customers,
            "top_products": [(row[0], row[1]) for row in top_products],
            "top_product_revenue": top_products[0][2] if top_products else 0,
            "channels": channels,
            "monthly_data": monthly_data,
            "shipped_orders": shipped_orders,
            "avg_delivery_time": avg_delivery_time,
            "main_regions": main_regions
        }
    except Exception as e:
        logger.error(f"Ошибка при получении данных для ежемесячного отчета: {str(e)}")
        return {
            "total_revenue": 0,
            "sales_count": 0,
            "avg_check": 0,
            "sales_dynamics": 0,
            "new_customers": 0,
            "top_products": [],
            "top_product_revenue": 0,
            "channels": [],
            "monthly_data": [],
            "shipped_orders": 0,
            "avg_delivery_time": "0",
            "main_regions": "Не указан"
        }


# Форма для недельного отчета
def create_weekly_word_report(start_date, end_date, data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Table for Logo and Company Info
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Left cell: Logo
    logo_cell = table.cell(0, 0)
    logo_cell.width = Inches(2.0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    # Замените 'path_to_logo.png' на актуальный путь к вашему логотипу
    logo_run.add_picture('/home/appuser/telegram-bot/logo.png', width=Inches(2.5))  # Увеличил размер логотипа до 1.5 дюйма

    # Right cell: Company Info
    company_cell = table.cell(0, 1)
    company_cell.width = Inches(4.5)
    company_info = (
        "ООО «Пример Компании»\n"
        "Рябиновая улица, 55с2, Москва, 121471\n"
        "Тел: +7 (495) 123-45-67, info@primercompany.ru\n"
        "ИНН: 1234567890, ОГРН: 1234567890123\n\n"
        "LLC «Example Company»\n"
        "Ryabinovaya street, 55c2, Moscow, 121471\n"
        "Phone: +7 (495) 123-45-67, info@primercompany.ru\n"
        "INN: 1234567890, OGRN: 1234567890123"
    )
    company_paragraph = company_cell.paragraphs[0]
    company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    company_run = company_paragraph.add_run(company_info)
    company_run.font.name = 'Times New Roman'
    company_run.font.size = Pt(10)
    company_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add some spacing after the table
    doc.add_paragraph()

    # Title: Report Title and Period
    title = doc.add_paragraph("ЕЖЕНЕДЕЛЬНЫЙ ОТЧЕТ ПО ПРОДАЖАМ")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    period = doc.add_paragraph(f"Отчетный период: {start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}")
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    period_run = period.runs[0]
    period_run.font.name = 'Times New Roman'
    period_run.font.size = Pt(12)
    period_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # Empty line

    # Section 1: Основные показатели
    heading = doc.add_paragraph("1. Общие показатели")
    heading.style = doc.styles['Heading 1']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    items = [
        f"Общая выручка: {data['total_revenue']:.2f} ₽",
        f"Количество продаж: {data['sales_count']}",
        f"Средний чек: {data['avg_check']:.2f} ₽",
        f"Динамика продаж: {data['sales_dynamics']:.2f}%",
        f"Количество новых покупателей: {data['new_customers']}"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(12)

    # Section 2: Анализ продаж
    heading = doc.add_paragraph("2. Анализ продаж")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    top_product = data['top_products'][0] if data['top_products'] else ("Не указан", 0)
    items = [
        f"Лучший продаваемый товар/услуга: {top_product[0]}",
        f"Количество проданных единиц: {top_product[1]}",
        f"Выручка от данного товара/услуги: {data.get('top_product_revenue', top_product[1] * 1000):.2f} ₽"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(12)

    # Section 3: Анализ каналов продаж
    heading = doc.add_paragraph("3. Анализ каналов продаж")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=len(data['channels']) + 1, cols=3)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)

    headers = ["Канал продаж", "Количество продаж", "Выручка (₽)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, (channel, sales, revenue) in enumerate(data['channels'], 1):
        table.cell(row_idx, 0).text = channel
        table.cell(row_idx, 1).text = str(sales)
        table.cell(row_idx, 2).text = f"{revenue:.2f}"
        for col_idx in range(3):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 4: Динамика продаж за неделю
    heading = doc.add_paragraph("4. Динамика продаж за неделю")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)
    table.columns[3].width = Inches(1.5)
    table.columns[4].width = Inches(2)

    headers = ["Дата", "Выручка (₽)", "Количество продаж", "Средний чек (₽)", "Изменение vs. прошлой недели (%)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    daily_data = data.get('daily_data', [])
    if not daily_data:
        current_date = start_date
        delta = (end_date - start_date).days + 1
        for row_idx in range(1, min(delta + 1, 7)):
            daily_revenue = data['total_revenue'] / delta if delta > 0 else 0
            daily_sales = data['sales_count'] / delta if delta > 0 else 0
            daily_avg_check = data['avg_check']
            change = data['sales_dynamics'] / delta if delta > 0 else 0
            table.cell(row_idx, 0).text = current_date.strftime('%d.%m.%Y')
            table.cell(row_idx, 1).text = f"{daily_revenue:.2f}"
            table.cell(row_idx, 2).text = f"{int(daily_sales)}"
            table.cell(row_idx, 3).text = f"{daily_avg_check:.2f}"
            table.cell(row_idx, 4).text = f"{change:.2f}"
            for col_idx in range(5):
                cell = table.cell(row_idx, col_idx)
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            current_date += timedelta(days=1)
    else:
        for row_idx, (date, revenue, sales, avg_check, change) in enumerate(daily_data[:6], 1):
            table.cell(row_idx, 0).text = date.strftime('%d.%m.%Y')
            table.cell(row_idx, 1).text = f"{revenue:.2f}"
            table.cell(row_idx, 2).text = f"{sales}"
            table.cell(row_idx, 3).text = f"{avg_check:.2f}"
            table.cell(row_idx, 4).text = f"{change:.2f}"
            for col_idx in range(5):
                cell = table.cell(row_idx, col_idx)
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    total_row_idx = min((end_date - start_date).days + 1, 7)
    table.cell(total_row_idx, 0).text = "Итого"
    table.cell(total_row_idx, 1).text = f"{data['total_revenue']:.2f}"
    table.cell(total_row_idx, 2).text = f"{data['sales_count']}"
    table.cell(total_row_idx, 3).text = f"{data['avg_check']:.2f}"
    table.cell(total_row_idx, 4).text = f"{data['sales_dynamics']:.2f}"
    for col_idx in range(5):
        cell = table.cell(total_row_idx, col_idx)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 5: Доставка
    heading = doc.add_paragraph("5. Доставка")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    items = [
        f"Количество заказов с доставкой: {data['shipped_orders']}",
        f"Среднее время доставки: {data['avg_delivery_time']}",
        f"Основные регионы доставки: {data['main_regions']}"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(12)

    # Footer: Date and Signature
    doc.add_paragraph()
    date = doc.add_paragraph(f'Дата составления отчета: {datetime.now().strftime("%d.%m.%Y")}')
    for run in date.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['Материально ответственное лицо', 'Аналитик продаж', '', 'подпись', '', 'расшифровка подписи']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
    table.cell(0, 5).text = 'А. В. Калинина'
    for p in table.cell(0, 5).paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_monthly_word_report(start_date, end_date, data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Table for Logo and Company Info
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Left cell: Logo
    logo_cell = table.cell(0, 0)
    logo_cell.width = Inches(2.0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    # Замените 'path_to_logo.png' на актуальный путь к вашему логотипу
    logo_run.add_picture('/home/appuser/telegram-bot/logo.png', width=Inches(2.5))  # Увеличил размер логотипа до 1.5 дюйма

    # Right cell: Company Info
    company_cell = table.cell(0, 1)
    company_cell.width = Inches(4.5)
    company_info = (
        "ООО «Пример Компании»\n"
        "Рябиновая улица, 55с2, Москва, 121471\n"
        "Тел: +7 (495) 123-45-67, info@primercompany.ru\n"
        "ИНН: 1234567890, ОГРН: 1234567890123\n\n"
        "LLC «Example Company»\n"
        "Ryabinovaya street, 55c2, Moscow, 121471\n"
        "Phone: +7 (495) 123-45-67, info@primercompany.ru\n"
        "INN: 1234567890, OGRN: 1234567890123"
    )
    company_paragraph = company_cell.paragraphs[0]
    company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    company_run = company_paragraph.add_run(company_info)
    company_run.font.name = 'Times New Roman'
    company_run.font.size = Pt(10)
    company_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add some spacing after the table
    doc.add_paragraph()

    # Title: Report Title and Period
    title = doc.add_paragraph("ЕЖЕМЕСЯЧНЫЙ ОТЧЕТ ПО ПРОДАЖАМ")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    period = doc.add_paragraph(f"Отчетный период: {start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}")
    period.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    period_run = period.runs[0]
    period_run.font.name = 'Times New Roman'
    period_run.font.size = Pt(12)
    period_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # Empty line

    # Section 1: Основные показатели
    heading = doc.add_paragraph("1. Общие показатели")
    heading.style = doc.styles['Heading 1']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    items = [
        f"Общая выручка: {data['total_revenue']:.2f} ₽",
        f"Количество продаж: {data['sales_count']}",
        f"Средний чек: {data['avg_check']:.2f} ₽",
        f"Динамика продаж: {data['sales_dynamics']:.2f}%",
        f"Количество новых покупателей: {data['new_customers']}"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(6)

    # Section 2: Анализ продаж
    heading = doc.add_paragraph("2. Анализ продаж")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    top_product = data['top_products'][0] if data['top_products'] else ("Не указан", 0)
    items = [
        f"Лучший продаваемый товар/услуга: {top_product[0]}",
        f"Количество проданных единиц: {top_product[1]}",
        f"Выручка от данного товара/услуги: {data.get('top_product_revenue', top_product[1] * 1000):.2f} ₽"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(6)

    # Section 3: Анализ каналов продаж
    heading = doc.add_paragraph("3. Анализ каналов продаж")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=len(data['channels']) + 1, cols=3)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)

    headers = ["Канал продаж", "Количество продаж", "Выручка (₽)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row_idx, (channel, sales, revenue) in enumerate(data['channels'], 1):
        table.cell(row_idx, 0).text = channel
        table.cell(row_idx, 1).text = str(sales)
        table.cell(row_idx, 2).text = f"{revenue:.2f}"
        for col_idx in range(3):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 4: Динамика продаж за месяц
    heading = doc.add_paragraph("4. Динамика продаж за месяц")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    table.autofit = True
    table.allow_autofit = True
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(1.5)
    table.columns[2].width = Inches(1.5)
    table.columns[3].width = Inches(1.5)
    table.columns[4].width = Inches(2)

    headers = ["Дата", "Выручка (₽)", "Количество продаж", "Средний чек (₽)", "Изменение vs. прошлый месяц (%)"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    monthly_data = data.get('monthly_data', [])
    if not monthly_data:
        prev_date = start_date - timedelta(days=30)
        table.cell(1, 0).text = f"Прошлый месяц {prev_date.strftime('%d.%m.%Y')}"
        table.cell(1, 1).text = f"{data['total_revenue'] * 0.9:.2f}"
        table.cell(1, 2).text = f"{int(data['sales_count'] * 0.9)}"
        table.cell(1, 3).text = f"{data['avg_check'] * 0.9:.2f}"
        table.cell(1, 4).text = "0.00"
        table.cell(2, 0).text = f"Нынешний месяц {start_date.strftime('%d.%m.%Y')}"
        table.cell(2, 1).text = f"{data['total_revenue']:.2f}"
        table.cell(2, 2).text = f"{data['sales_count']}"
        table.cell(2, 3).text = f"{data['avg_check']:.2f}"
        table.cell(2, 4).text = f"{data['sales_dynamics']:.2f}"
    else:
        for row_idx, (date, revenue, sales, avg_check, change) in enumerate(monthly_data[:2], 1):
            table.cell(row_idx, 0).text = date.strftime('%d.%m.%Y')
            table.cell(row_idx, 1).text = f"{revenue:.2f}"
            table.cell(row_idx, 2).text = f"{sales}"
            table.cell(row_idx, 3).text = f"{avg_check:.2f}"
            table.cell(row_idx, 4).text = f"{change:.2f}"

    for row_idx in range(1, 3):
        for col_idx in range(5):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Section 5: Доставка
    heading = doc.add_paragraph("5. Доставка")
    heading.style = doc.styles['Heading 2']
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    items = [
        f"Количество заказов с доставкой: {data['shipped_orders']}",
        f"Среднее время доставки: {data['avg_delivery_time']}",
        f"Основные регионы доставки: {data['main_regions']}"
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.style.font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(6)

    # Footer: Date and Signature
    doc.add_paragraph()
    date = doc.add_paragraph(f'Дата составления отчета: {datetime.now().strftime("%d.%m.%Y")}')
    for run in date.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['Материально ответственное лицо', 'Аналитик продаж', '', '', '', 'расшифровка подписи']
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
    table.cell(0, 5).text = 'А. В. Калинина'
    for p in table.cell(0, 5).paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Хранилище состояния пользователя
user_state = {}
dp = Dispatcher()   
@dp.message(Command("start"))
async def send_welcome(message: types.Message):
    user_id = message.from_user.id
    logger.info(f"Получена команда /start от user_id: {user_id}, текущий user_state: {user_state.get(user_id)}")
    await message.answer(
        "<b>👋 Привет! Я BI Mate — твой умный помощник.</b>\n"
        "<b>Строю отчеты, интерактивные графики и дашборды, анализирую товары и продажи.</b>\n\n"
        "Давай упростим твою работу — с чего начнем?\n"
        "<i>Если у вас есть вопросы, напишите @foolforu1</i>", 
        reply_markup=main_menu,
        parse_mode="HTML" 
        
    )

# Обработчик кнопки "Графики"
@dp.message(lambda message: message.text == "Графики")
async def show_graph_menu(message: types.Message):
    try:
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Динамика выручки", callback_data="graph_type_sales_dynamics")],
            [InlineKeyboardButton(text="Категории товаров", callback_data="graph_type_category_sales")],
            [InlineKeyboardButton(text="Выручка по городам", callback_data="graph_type_city_revenue")],
            [InlineKeyboardButton(text="Методы оплаты", callback_data="graph_type_payment_methods")],
            [InlineKeyboardButton(text="Пол покупателей", callback_data="graph_type_gender_stats")],
            [InlineKeyboardButton(text="Топ-10 товаров", callback_data="graph_type_top_goods")],
            [InlineKeyboardButton(text="Динамика заказов", callback_data="graph_type_order_dynamics")],
            [InlineKeyboardButton(text="Топ-15 брендов", callback_data="graph_type_top_brend")]
        ])
        await message.answer("Выберите тип графика:", reply_markup=keyboard)
    except Exception as e:
        logger.error(f"Ошибка в show_graph_menu: {str(e)}")
        await message.answer("Произошла ошибка при отображении меню графиков. Попробуйте снова.")

# Обработчик кнопки "Отчеты"
@dp.message(lambda message: message.text == "Отчеты")
async def show_report_menu(message: types.Message):
    try:
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Дашборд", callback_data="report_type_Дашборд")],
            [InlineKeyboardButton(text="Еженедельный", callback_data="report_type_Еженедельный")],
            [InlineKeyboardButton(text="Ежемесячный", callback_data="report_type_Ежемесячный")]
        ])
        await message.answer("Выберите тип отчета:", reply_markup=keyboard)
    except Exception as e:
        logger.error(f"Ошибка в show_report_menu: {str(e)}")
        await message.answer("Произошла ошибка при отображении меню отчетов. Попробуйте снова.")

#Обработчик кнопки "Анализ"
@dp.message(lambda message: message.text == "Анализ товара")
async def show_product_analysis(message: types.Message):
    try:
        web_url = "http://91.135.156.17/products"  # После настройки Nginx
        # Если Nginx не настроен, используйте: web_url = "http://91.135.156.17:5000"
        await message.answer(
            f"Откройте эту ссылку в браузере для интерактивного анализа товаров:\n{web_url}",
            reply_markup=main_menu
        )
    except Exception as e:
        logger.error(f"Ошибка в show_product_analysis: {str(e)}")
        await message.answer("Произошла ошибка при запуске анализа.")

@dp.message(lambda message: message.text == "Анализ продаж")
async def show_sales_analysis(message: types.Message):
    try:
        web_url = "http://91.135.156.17/sales"  # URL для анализа продаж
        await message.answer(
            f"Откройте эту ссылку в браузере для интерактивного анализа продаж:\n{web_url}",
            reply_markup=main_menu
        )
    except Exception as e:
        logger.error(f"Ошибка в show_sales_analysis: {str(e)}")
        await message.answer("Произошла ошибка при запуске анализа продаж.")

@dp.callback_query(lambda c: c.data.startswith("graph_type_"))
async def process_graph_type(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        graph_type = callback.data.split("_", 2)[2]
        user_state[user_id] = {
            "type": "graph",
            "graph_type": graph_type,
            "history": ["graph_menu"]  # Начальный шаг - меню графиков
        }
        
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Год", callback_data="graph_period_year")],
            [InlineKeyboardButton(text="Полгода", callback_data="graph_period_halfyear")],
            [InlineKeyboardButton(text="Кварталы", callback_data="graph_period_quarter")],
            [InlineKeyboardButton(text="Месяц", callback_data="graph_period_month")],
            [InlineKeyboardButton(text="Недели", callback_data="graph_period_week")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_menu")]
        ])
        await callback.message.edit_text("Выберите период для графика:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_type: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

# Обработчик выбора типа отчета
@dp.callback_query(lambda c: c.data.startswith("report_type_"))
async def process_report_type(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        report_type = callback.data.split("_", 2)[2]
        user_state[user_id] = {"type": "report", "report_type": report_type, "history": ["report_menu"]}
        
        if report_type == "Дашборд":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="По году", callback_data="dashboard_period_year")],
                [InlineKeyboardButton(text="По месяцам", callback_data="dashboard_period_month")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
            ])
            await callback.message.edit_text("Выберите период для дашборда:", reply_markup=keyboard)
        else:
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="report_year_prev"),
                 InlineKeyboardButton(text=str(datetime.now().year), callback_data="report_year_select"),
                 InlineKeyboardButton(text="➡️", callback_data="report_year_next")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
            ])
            await callback.message.edit_text(f"Выберите год для отчета:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_report_type: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

# Обработчик выбора периода для дашборда
@dp.callback_query(lambda c: c.data.startswith("dashboard_period_"))
async def process_dashboard_period(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        period = callback.data.split("_")[2]
        user_state[user_id]["period"] = period
        current_year = datetime.now().year
        user_state[user_id]["year"] = current_year
        user_state[user_id]["history"].append("dashboard_period")  # Сохраняем шаг
        
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="report_year_prev"),
             InlineKeyboardButton(text=str(current_year), callback_data="report_year_select"),
             InlineKeyboardButton(text="➡️", callback_data="report_year_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
        ])
        await callback.message.edit_text(f"Выберите год для дашборда:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_dashboard_period: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

#Обработчик кнопка Назад
@dp.callback_query(lambda c: c.data.startswith("back_to_"))
async def process_back(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        if user_id not in user_state:
            await callback.message.edit_text("Состояние утеряно. Начните заново.", reply_markup=None)
            await callback.message.answer("Выберите действие:", reply_markup=main_menu)
            return

        # Получаем текущий шаг и текст сообщения
        current_text = callback.message.text
        history = user_state[user_id].get("history", [])
        if not history:
            await callback.message.edit_text("История шагов пуста. Вернитесь в главное меню.", reply_markup=None)
            await callback.message.answer("Выберите действие:", reply_markup=main_menu)
            if user_id in user_state:
                del user_state[user_id]
            return

        # Удаляем текущий шаг из истории
        previous_step = history.pop()

        # Проверяем, есть ли еще шаги в истории, чтобы определить следующий шаг
        next_step = history[-1] if history else None

        # Логика возврата для отчетов
        if previous_step == "report_menu" or next_step == "report_menu":
            if current_text != "Выберите тип отчета:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="Дашборд", callback_data="report_type_Дашборд")],
                    [InlineKeyboardButton(text="Еженедельный", callback_data="report_type_Еженедельный")],
                    [InlineKeyboardButton(text="Ежемесячный", callback_data="report_type_Ежемесячный")]
                ])
                await callback.message.edit_text("Выберите тип отчета:", reply_markup=keyboard)
                user_state[user_id]["history"] = ["report_menu"]

        elif previous_step == "report_year" or next_step == "report_year":
            year = user_state[user_id]["year"]
            if current_text != "Выберите год для отчета:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="report_year_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="report_year_select"),
                     InlineKeyboardButton(text="➡️", callback_data="report_year_next")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
                ])
                await callback.message.edit_text(f"Выберите год для отчета:", reply_markup=keyboard)
                if "report_year" not in history:
                    history.append("report_year")

        elif previous_step == "report_month" or next_step == "report_month":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите месяц для отчета ({year}):":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="Янв", callback_data="report_month_1"),
                     InlineKeyboardButton(text="Фев", callback_data="report_month_2"),
                     InlineKeyboardButton(text="Мар", callback_data="report_month_3")],
                    [InlineKeyboardButton(text="Апр", callback_data="report_month_4"),
                     InlineKeyboardButton(text="Май", callback_data="report_month_5"),
                     InlineKeyboardButton(text="Июн", callback_data="report_month_6")],
                    [InlineKeyboardButton(text="Июл", callback_data="report_month_7"),
                     InlineKeyboardButton(text="Авг", callback_data="report_month_8"),
                     InlineKeyboardButton(text="Сен", callback_data="report_month_9")],
                    [InlineKeyboardButton(text="Окт", callback_data="report_month_10"),
                     InlineKeyboardButton(text="Ноя", callback_data="report_month_11"),
                     InlineKeyboardButton(text="Дек", callback_data="report_month_12")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_report_year")]
                ])
                await callback.message.edit_text(f"Выберите месяц для отчета ({year}):", reply_markup=keyboard)
                if "report_month" not in history:
                    history.append("report_month")

        elif previous_step == "report_week_selection" or next_step == "report_week_selection":
            year = user_state[user_id]["year"]
            month = user_state[user_id]["month"]
            weeks = user_state[user_id]["weeks"]
            if current_text != f"Выберите неделю для отчета ({year}-{month:02d}):":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(
                        text=f"Неделя {i+1} ({datetime(year, month, start).strftime('%d.%m')} - {datetime(year, month, end).strftime('%d.%m')})",
                        callback_data=f"report_week_select_{i}"
                    )] for i, (start, end) in enumerate(weeks)
                ] + [[InlineKeyboardButton(text="Назад", callback_data="back_to_report_month")]])
                await callback.message.edit_text(f"Выберите неделю для отчета ({year}-{month:02d}):", reply_markup=keyboard)
                if "report_week_selection" not in history:
                    history.append("report_week_selection")

        # Логика возврата для дашборда
        elif previous_step == "dashboard_period" or next_step == "dashboard_period":
            if current_text != "Выберите период для дашборда:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="По году", callback_data="dashboard_period_year")],
                    [InlineKeyboardButton(text="По месяцам", callback_data="dashboard_period_month")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
                ])
                await callback.message.edit_text("Выберите период для дашборда:", reply_markup=keyboard)
                if "dashboard_period" not in history:
                    history.append("dashboard_period")

        # Логика возврата для графиков
        elif previous_step == "graph_type" or next_step == "graph_type":
            if current_text != "Выберите тип графика:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="Динамика выручки", callback_data="graph_type_sales_dynamics")],
                    [InlineKeyboardButton(text="Категории товаров", callback_data="graph_type_category_sales")],
                    [InlineKeyboardButton(text="Выручка по городам", callback_data="graph_type_city_revenue")],
                    [InlineKeyboardButton(text="Методы оплаты", callback_data="graph_type_payment_methods")],
                    [InlineKeyboardButton(text="Пол покупателей", callback_data="graph_type_gender_stats")],
                    [InlineKeyboardButton(text="Топ-10 товаров", callback_data="graph_type_top_goods")],
                    [InlineKeyboardButton(text="Динамика заказов", callback_data="graph_type_order_dynamics")],
                    [InlineKeyboardButton(text="Топ-15 брендов", callback_data="graph_type_top_brend")]
                ])
                await callback.message.edit_text("Выберите тип графика:", reply_markup=keyboard)
                if "graph_type" not in history:
                    history.append("graph_type")

        elif previous_step == "graph_period" or next_step == "graph_period":
            if current_text != "Выберите период для графика:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="По году", callback_data="graph_period_year")],
                    [InlineKeyboardButton(text="По полугодиям", callback_data="graph_period_halfyear")],
                    [InlineKeyboardButton(text="По кварталам", callback_data="graph_period_quarter")],
                    [InlineKeyboardButton(text="По месяцам", callback_data="graph_period_month")],
                    [InlineKeyboardButton(text="По неделям", callback_data="graph_period_week")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_menu")]
                ])
                await callback.message.edit_text("Выберите период для графика:", reply_markup=keyboard)
                if "graph_period" not in history:
                    history.append("graph_period")

        elif previous_step == "graph_month" or next_step == "graph_month":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите год для месяца:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="graph_month_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="graph_month_select"),
                     InlineKeyboardButton(text="➡️", callback_data="graph_month_next")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
                ])
                await callback.message.edit_text(f"Выберите год для месяца:", reply_markup=keyboard)
                if "graph_month" not in history:
                    history.append("graph_month")

        elif previous_step == "graph_month_select" or next_step == "graph_month_select":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите год для месяца:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="graph_month_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="graph_month_select"),
                     InlineKeyboardButton(text="➡️", callback_data="graph_month_next")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
                ])
                await callback.message.edit_text(f"Выберите год для месяца:", reply_markup=keyboard)
                if "graph_month_select" not in history:
                    history.append("graph_month_select")

        elif previous_step == "graph_week_year" or next_step == "graph_week_year":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите год для недельного графика:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="graph_week_year_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="graph_week_year_select"),
                     InlineKeyboardButton(text="➡️", callback_data="graph_week_year_next")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
                ])
                await callback.message.edit_text(f"Выберите год для недельного графика:", reply_markup=keyboard)
                if "graph_week_year" not in history:
                    history.append("graph_week_year")

        elif previous_step == "graph_week_month" or next_step == "graph_week_month":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите месяц для недельного графика ({year}):":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="Янв", callback_data="graph_week_month_1"),
                     InlineKeyboardButton(text="Фев", callback_data="graph_week_month_2"),
                     InlineKeyboardButton(text="Мар", callback_data="graph_week_month_3")],
                    [InlineKeyboardButton(text="Апр", callback_data="graph_week_month_4"),
                     InlineKeyboardButton(text="Май", callback_data="graph_week_month_5"),
                     InlineKeyboardButton(text="Июн", callback_data="graph_week_month_6")],
                    [InlineKeyboardButton(text="Июл", callback_data="graph_week_month_7"),
                     InlineKeyboardButton(text="Авг", callback_data="graph_week_month_8"),
                     InlineKeyboardButton(text="Сен", callback_data="graph_week_month_9")],
                    [InlineKeyboardButton(text="Окт", callback_data="graph_week_month_10"),
                     InlineKeyboardButton(text="Ноя", callback_data="graph_week_month_11"),
                     InlineKeyboardButton(text="Дек", callback_data="graph_week_month_12")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_week_year")]
                ])
                await callback.message.edit_text(f"Выберите месяц для недельного графика ({year}):", reply_markup=keyboard)
                if "graph_week_month" not in history:
                    history.append("graph_week_month")

        elif previous_step == "graph_week_selection" or next_step == "graph_week_selection":
            year = user_state[user_id]["year"]
            month = user_state[user_id]["month"]
            weeks = user_state[user_id]["weeks"]
            if current_text != f"Выберите неделю для графика ({year}-{month:02d}):":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(
                        text=f"Неделя {i+1} ({datetime(year, month, start).strftime('%d.%m')} - {datetime(year, month, end).strftime('%d.%m')})",
                        callback_data=f"graph_week_select_{i}"
                    )] for i, (start, end) in enumerate(weeks)
                ] + [[InlineKeyboardButton(text="Назад", callback_data="back_to_graph_week_month")]])
                await callback.message.edit_text(f"Выберите неделю для графика ({year}-{month:02d}):", reply_markup=keyboard)
                if "graph_week_selection" not in history:
                    history.append("graph_week_selection")

        elif previous_step == "graph_halfyear" or next_step == "graph_halfyear":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите полугодие для графика:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="graph_halfyear_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="graph_halfyear_select"),
                     InlineKeyboardButton(text="➡️", callback_data="graph_halfyear_next")],
                    [InlineKeyboardButton(text="1-е полугодие", callback_data="graph_halfyear_1"),
                     InlineKeyboardButton(text="2-е полугодие", callback_data="graph_halfyear_2")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
                ])
                await callback.message.edit_text(f"Выберите полугодие для графика:", reply_markup=keyboard)
                if "graph_halfyear" not in history:
                    history.append("graph_halfyear")

        elif previous_step == "graph_quarter" or next_step == "graph_quarter":
            year = user_state[user_id]["year"]
            if current_text != f"Выберите квартал для графика:":
                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(text="⬅️", callback_data="graph_quarter_prev"),
                     InlineKeyboardButton(text=str(year), callback_data="graph_quarter_select"),
                     InlineKeyboardButton(text="➡️", callback_data="graph_quarter_next")],
                    [InlineKeyboardButton(text="1 кв", callback_data="graph_quarter_1"),
                     InlineKeyboardButton(text="2 кв", callback_data="graph_quarter_2"),
                     InlineKeyboardButton(text="3 кв", callback_data="graph_quarter_3"),
                     InlineKeyboardButton(text="4 кв", callback_data="graph_quarter_4")],
                    [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
                ])
                await callback.message.edit_text(f"Выберите квартал для графика:", reply_markup=keyboard)
                if "graph_quarter" not in history:
                    history.append("graph_quarter")

        # Обработка возврата в главное меню
        elif previous_step == "main_menu":
            await callback.message.edit_text("Вы вернулись в главное меню.", reply_markup=None)
            await callback.message.answer("Выберите действие:", reply_markup=main_menu)
            if user_id in user_state:
                del user_state[user_id]

        else:
            # Вместо сброса состояния пытаемся вернуться к выбору типа отчета
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="Дашборд", callback_data="report_type_Дашборд")],
                [InlineKeyboardButton(text="Еженедельный", callback_data="report_type_Еженедельный")],
                [InlineKeyboardButton(text="Ежемесячный", callback_data="report_type_Ежемесячный")]
            ])
            await callback.message.edit_text("Выберите тип отчета:", reply_markup=keyboard)
            user_state[user_id]["history"] = ["report_menu"]

        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_back: {str(e)}")
        await callback.message.edit_text("Произошла ошибка при возврате назад. Попробуйте снова.", reply_markup=None)
        await callback.message.answer("Выберите действие:", reply_markup=main_menu)
        if user_id in user_state:
            del user_state[user_id]


# Обработчик Период для графиков
@dp.callback_query(lambda c: c.data.startswith("graph_period_"))
async def process_graph_period(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        period = callback.data.split("_")[2]
        current_year = datetime.now().year
        user_state[user_id]["period"] = period
        user_state[user_id]["year"] = current_year

        if period == "year":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="graph_year_prev"),
                 InlineKeyboardButton(text=str(current_year), callback_data="graph_year_select"),
                 InlineKeyboardButton(text="➡️", callback_data="graph_year_next")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
            ])
            await callback.message.edit_text(f"Выберите год для графика:", reply_markup=keyboard)
            user_state[user_id]["history"].append("graph_period")
        elif period == "halfyear":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="graph_halfyear_prev"),
                 InlineKeyboardButton(text=str(current_year), callback_data="graph_halfyear_select"),
                 InlineKeyboardButton(text="➡️", callback_data="graph_halfyear_next")],
                [InlineKeyboardButton(text="1-е полугодие", callback_data="graph_halfyear_1"),
                 InlineKeyboardButton(text="2-е полугодие", callback_data="graph_halfyear_2")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
            ])
            await callback.message.edit_text(f"Выберите полугодие для графика:", reply_markup=keyboard)
            user_state[user_id]["history"].append("graph_period")
        elif period == "quarter":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="graph_quarter_prev"),
                 InlineKeyboardButton(text=str(current_year), callback_data="graph_quarter_select"),
                 InlineKeyboardButton(text="➡️", callback_data="graph_quarter_next")],
                [InlineKeyboardButton(text="1 кв", callback_data="graph_quarter_1"),
                 InlineKeyboardButton(text="2 кв", callback_data="graph_quarter_2"),
                 InlineKeyboardButton(text="3 кв", callback_data="graph_quarter_3"),
                 InlineKeyboardButton(text="4 кв", callback_data="graph_quarter_4")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
            ])
            await callback.message.edit_text(f"Выберите квартал для графика:", reply_markup=keyboard)
            user_state[user_id]["history"].append("graph_period")
        elif period == "month":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="graph_month_prev"),
                 InlineKeyboardButton(text=str(current_year), callback_data="graph_month_select"),
                 InlineKeyboardButton(text="➡️", callback_data="graph_month_next")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
            ])
            await callback.message.edit_text(f"Выберите год для месяца:", reply_markup=keyboard)
            user_state[user_id]["history"].append("graph_month")
        elif period == "week":
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="⬅️", callback_data="graph_week_year_prev"),
                 InlineKeyboardButton(text=str(current_year), callback_data="graph_week_year_select"),
                 InlineKeyboardButton(text="➡️", callback_data="graph_week_year_next")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
            ])
            await callback.message.edit_text(f"Выберите год для недельного графика:", reply_markup=keyboard)
            user_state[user_id]["history"].append("graph_week_year")
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_period: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=main_menu)

#Обработчик для графиков выбора года для графиков
@dp.callback_query(lambda c: c.data.startswith("graph_year_"))
async def process_graph_year(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id]["year"]

        if action == "prev":
            user_state[user_id]["year"] = year - 1
        elif action == "next":
            user_state[user_id]["year"] = year + 1
        elif action == "select":
            start_date = datetime(year, 1, 1)
            end_date = datetime(year, 12, 31)
            graph_type = user_state[user_id]["graph_type"]
            graph_buffer, error = create_graph(graph_type, start_date, end_date)
            if error:
                await callback.message.edit_text(error, reply_markup=None)
                del user_state[user_id]
                return
            await callback.message.delete()
            await callback.message.answer_photo(
                photo=BufferedInputFile(graph_buffer.read(), filename=f"{graph_type}_{year}.png"),
                caption=f"График: {graph_type} за {year} год",
                reply_markup=main_menu
            )
            del user_state[user_id]
            return

        year = user_state[user_id]["year"]
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="graph_year_prev"),
             InlineKeyboardButton(text=str(year), callback_data="graph_year_select"),
             InlineKeyboardButton(text="➡️", callback_data="graph_year_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_menu")]
        ])
        await callback.message.edit_text(f"Выберите год для графика:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_year: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

#Обработчик для графиков выбора полугодия
@dp.callback_query(lambda c: c.data.startswith("graph_halfyear_"))
async def process_graph_halfyear(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id]["year"]

        if action == "prev":
            user_state[user_id]["year"] = year - 1
        elif action == "next":
            user_state[user_id]["year"] = year + 1
        elif action in ["1", "2"]:
            halfyear = int(action)
            if halfyear == 1:
                start_date = datetime(year, 1, 1)
                end_date = datetime(year, 6, 30)
            else:
                start_date = datetime(year, 7, 1)
                end_date = datetime(year, 12, 31)
            graph_type = user_state[user_id]["graph_type"]
            graph_buffer, error = create_graph(graph_type, start_date, end_date)
            if error:
                await callback.message.edit_text(error, reply_markup=None)
                del user_state[user_id]
                return
            await callback.message.delete()
            await callback.message.answer_photo(
                photo=BufferedInputFile(graph_buffer.read(), filename=f"{graph_type}_halfyear_{halfyear}_{year}.png"),
                caption=f"График: {graph_type} за {halfyear}-е полугодие {year}",
                reply_markup=main_menu
            )
            del user_state[user_id]
            return

        year = user_state[user_id]["year"]
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="graph_halfyear_prev"),
             InlineKeyboardButton(text=str(year), callback_data="graph_halfyear_select"),
             InlineKeyboardButton(text="➡️", callback_data="graph_halfyear_next")],
            [InlineKeyboardButton(text="1-е полугодие", callback_data="graph_halfyear_1"),
             InlineKeyboardButton(text="2-е полугодие", callback_data="graph_halfyear_2")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_menu")]
        ])
        await callback.message.edit_text(f"Выберите полугодие для графика:", reply_markup=keyboard)
        user_state[user_id]["history"].append("graph_halfyear")  # Сохраняем шаг
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_halfyear: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

#Обработчик для графиков выбора квартала
@dp.callback_query(lambda c: c.data.startswith("graph_quarter_"))
async def process_graph_quarter(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id]["year"]

        if action == "prev":
            user_state[user_id]["year"] = year - 1
        elif action == "next":
            user_state[user_id]["year"] = year + 1
        elif action in ["1", "2", "3", "4"]:
            quarter = int(action)
            if quarter == 1:
                start_date = datetime(year, 1, 1)
                end_date = datetime(year, 3, 31)
            elif quarter == 2:
                start_date = datetime(year, 4, 1)
                end_date = datetime(year, 6, 30)
            elif quarter == 3:
                start_date = datetime(year, 7, 1)
                end_date = datetime(year, 9, 30)
            else:
                start_date = datetime(year, 10, 1)
                end_date = datetime(year, 12, 31)
            graph_type = user_state[user_id]["graph_type"]
            graph_buffer, error = create_graph(graph_type, start_date, end_date)
            if error:
                await callback.message.edit_text(error, reply_markup=None)
                del user_state[user_id]
                return
            await callback.message.delete()
            await callback.message.answer_photo(
                photo=BufferedInputFile(graph_buffer.read(), filename=f"{graph_type}_quarter_{quarter}_{year}.png"),
                caption=f"График: {graph_type} за {quarter}-й квартал {year}",
                reply_markup=main_menu
            )
            del user_state[user_id]
            return

        year = user_state[user_id]["year"]
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="graph_quarter_prev"),
             InlineKeyboardButton(text=str(year), callback_data="graph_quarter_select"),
             InlineKeyboardButton(text="➡️", callback_data="graph_quarter_next")],
            [InlineKeyboardButton(text="1 кв", callback_data="graph_quarter_1"),
             InlineKeyboardButton(text="2 кв", callback_data="graph_quarter_2"),
             InlineKeyboardButton(text="3 кв", callback_data="graph_quarter_3"),
             InlineKeyboardButton(text="4 кв", callback_data="graph_quarter_4")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_menu")]
        ])
        await callback.message.edit_text(f"Выберите квартал для графика:", reply_markup=keyboard)
        user_state[user_id]["history"].append("graph_quarter")  # Сохраняем шаг
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_quarter: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)

#Обработчик для графиков выбора месяца
@dp.callback_query(lambda c: c.data.startswith("graph_month_"))
async def process_graph_month(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id]["year"]

        if action == "prev":
            year -= 1
        elif action == "next":
            year += 1
        elif action == "select":
            user_state[user_id]["history"].append("graph_month_select")
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="Янв", callback_data="graph_month_1"),
                 InlineKeyboardButton(text="Фев", callback_data="graph_month_2"),
                 InlineKeyboardButton(text="Мар", callback_data="graph_month_3")],
                [InlineKeyboardButton(text="Апр", callback_data="graph_month_4"),
                 InlineKeyboardButton(text="Май", callback_data="graph_month_5"),
                 InlineKeyboardButton(text="Июн", callback_data="graph_month_6")],
                [InlineKeyboardButton(text="Июл", callback_data="graph_month_7"),
                 InlineKeyboardButton(text="Авг", callback_data="graph_month_8"),
                 InlineKeyboardButton(text="Сен", callback_data="graph_month_9")],
                [InlineKeyboardButton(text="Окт", callback_data="graph_month_10"),
                 InlineKeyboardButton(text="Ноя", callback_data="graph_month_11"),
                 InlineKeyboardButton(text="Дек", callback_data="graph_month_12")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_month")]
            ])
            await callback.message.edit_text(f"Выберите месяц для графика ({year}):", reply_markup=keyboard)
            await callback.answer()
            return

        if action.isdigit():
            month = int(action)
            start_date = datetime(year, month, 1).date()
            end_date = datetime(year, month, monthrange(year, month)[1]).date()
            graph_type = user_state[user_id]["graph_type"]
            graph_buffer, error_message = create_graph(graph_type, start_date, end_date)
            if graph_buffer:
                await bot.send_photo(
                    chat_id=callback.message.chat.id,
                    photo=BufferedInputFile(graph_buffer.getvalue(), filename=f"{graph_type}.png"),
                    reply_markup=main_menu
                )
            else:
                await callback.message.edit_text(error_message or "Нет данных для построения графика.", reply_markup=main_menu)
            del user_state[user_id]
            await callback.message.delete()
            return

        user_state[user_id]["year"] = year
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="graph_month_prev"),
             InlineKeyboardButton(text=str(year), callback_data="graph_month_select"),
             InlineKeyboardButton(text="➡️", callback_data="graph_month_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
        ])
        await callback.message.edit_text(f"Выберите год для месяца:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_month: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=main_menu)

#Обработчик для графика выбора года в недельном периоде
@dp.callback_query(lambda c: c.data.startswith("graph_week_year_"))
async def process_graph_week_year(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        logger.info(f"Вызов process_graph_week_year для user_id={user_id}, callback_data={callback.data}")
        action = callback.data.split("_")[3]
        year = user_state[user_id]["year"]
        logger.info(f"Текущий год: {year}, действие: {action}")

        if action == "prev":
            year -= 1
        elif action == "next":
            year += 1
        elif action == "select":
            logger.info(f"Переход к выбору месяца для user_id={user_id}")
            user_state[user_id]["history"].append("graph_week_month")
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="Янв", callback_data="graph_week_month_1"),
                 InlineKeyboardButton(text="Фев", callback_data="graph_week_month_2"),
                 InlineKeyboardButton(text="Мар", callback_data="graph_week_month_3")],
                [InlineKeyboardButton(text="Апр", callback_data="graph_week_month_4"),
                 InlineKeyboardButton(text="Май", callback_data="graph_week_month_5"),
                 InlineKeyboardButton(text="Июн", callback_data="graph_week_month_6")],
                [InlineKeyboardButton(text="Июл", callback_data="graph_week_month_7"),
                 InlineKeyboardButton(text="Авг", callback_data="graph_week_month_8"),
                 InlineKeyboardButton(text="Сен", callback_data="graph_week_month_9")],
                [InlineKeyboardButton(text="Окт", callback_data="graph_week_month_10"),
                 InlineKeyboardButton(text="Ноя", callback_data="graph_week_month_11"),
                 InlineKeyboardButton(text="Дек", callback_data="graph_week_month_12")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_week_year")]
            ])
            await callback.message.edit_text(f"Выберите месяц для недельного графика ({year}):", reply_markup=keyboard)
            await callback.answer()
            return

        user_state[user_id]["year"] = year
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="graph_week_year_prev"),
             InlineKeyboardButton(text=str(year), callback_data="graph_week_year_select"),
             InlineKeyboardButton(text="➡️", callback_data="graph_week_year_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_graph_period")]
        ])
        await callback.message.edit_text(f"Выберите год для недельного графика:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_week_year: {str(e)}")
        await callback.message.delete()
        await callback.message.answer("Произошла ошибка. Попробуйте снова.", reply_markup=main_menu)
        if user_id in user_state:
            del user_state[user_id]

#Обработчик для графика выбора месяца и недели,
@dp.callback_query(lambda c: c.data.startswith("graph_week_month_"))
async def process_graph_week_month(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        month = int(callback.data.split("_")[3])
        year = user_state[user_id]["year"]

        weeks = []
        cal = calendar.Calendar(firstweekday=0)
        for week in cal.monthdayscalendar(year, month):
            if week[0] != 0:
                start_day = week[0]
                end_day = week[6] if week[6] != 0 else monthrange(year, month)[1]
                if start_day <= monthrange(year, month)[1] and end_day <= monthrange(year, month)[1]:
                    weeks.append((start_day, end_day))

        if not weeks:
            await callback.message.edit_text("Нет доступных недель для выбранного месяца.", reply_markup=main_menu)
            del user_state[user_id]
            await callback.message.delete()
            return

        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(
                text=f"Неделя {i+1} ({datetime(year, month, start).strftime('%d.%m')} - {datetime(year, month, end).strftime('%d.%m')})",
                callback_data=f"graph_week_select_{i}"
            )] for i, (start, end) in enumerate(weeks)
        ] + [[InlineKeyboardButton(text="Назад", callback_data="back_to_graph_week_month")]])
        user_state[user_id]["month"] = month
        user_state[user_id]["weeks"] = weeks
        user_state[user_id]["history"].append("graph_week_selection")
        await callback.message.edit_text(f"Выберите неделю для графика ({year}-{month:02d}):", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_week_month: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=main_menu)

#Обработчик
@dp.callback_query(lambda c: c.data.startswith("graph_week_select_"))
async def process_graph_week_select(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        week_idx = int(callback.data.split("_")[3])
        year = user_state[user_id]["year"]
        month = user_state[user_id]["month"]
        start_day, end_day = user_state[user_id]["weeks"][week_idx]
        start_date = datetime(year, month, start_day).date()
        end_date = datetime(year, month, end_day).date()
        
        graph_type = user_state[user_id]["graph_type"]
        graph_buffer, error_message = create_graph(graph_type, start_date, end_date)
        if graph_buffer:
            await callback.message.delete()
            await bot.send_photo(
                chat_id=callback.message.chat.id,
                photo=BufferedInputFile(graph_buffer.getvalue(), filename=f"{graph_type}.png"),
                reply_markup=main_menu
            )
        else:
            await callback.message.delete()
            await callback.message.answer(
                error_message or "Нет данных для построения графика за выбранный период.",
                reply_markup=main_menu
            )
        if user_id in user_state:
            del user_state[user_id]
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_graph_week_select: {str(e)}")
        await callback.message.delete()
        await callback.message.answer("Произошла ошибка при построении графика. Попробуйте снова.", reply_markup=main_menu)
        if user_id in user_state:
            del user_state[user_id]
                
#Обработчик для отчета (год)        
@dp.callback_query(lambda c: c.data.startswith("report_year_"))
async def process_report_year(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id].get("year", datetime.now().year)
        
        if action == "prev":
            year -= 1
        elif action == "next":
            year += 1
        elif action == "select":
            user_state[user_id]["year"] = year
            report_type = user_state[user_id]["report_type"]
            
            if report_type == "Дашборд" and user_state[user_id].get("period") == "year":
                start_date = datetime(year, 1, 1).date()
                end_date = datetime(year, 12, 31).date()
                graph_buffer = create_dashboard(start_date, end_date)
                if not graph_buffer:
                    await callback.message.edit_text("Нет данных для построения дашборда.", reply_markup=main_menu)
                    del user_state[user_id]
                    await callback.message.delete()
                    return
                data = get_dashboard_table_data(start_date, end_date)
                column_names = ["Показатель", "Значение"]
                pdf_buffer = create_pdf("Дашборд", graph_buffer, data, column_names, start_date, end_date)
                if pdf_buffer:
                    filename = f"Дашборд_{start_date.strftime('%Y')}.pdf"
                    await bot.send_document(
                        chat_id=callback.message.chat.id,
                        document=BufferedInputFile(pdf_buffer.getvalue(), filename=filename),
                        reply_markup=main_menu
                    )
                else:
                    await callback.message.edit_text("Ошибка при создании дашборда.", reply_markup=main_menu)
                del user_state[user_id]
                await callback.message.delete()
                return
            
            user_state[user_id]["history"].append("report_year")  # Сохраняем шаг
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="Янв", callback_data="report_month_1"),
                 InlineKeyboardButton(text="Фев", callback_data="report_month_2"),
                 InlineKeyboardButton(text="Мар", callback_data="report_month_3")],
                [InlineKeyboardButton(text="Апр", callback_data="report_month_4"),
                 InlineKeyboardButton(text="Май", callback_data="report_month_5"),
                 InlineKeyboardButton(text="Июн", callback_data="report_month_6")],
                [InlineKeyboardButton(text="Июл", callback_data="report_month_7"),
                 InlineKeyboardButton(text="Авг", callback_data="report_month_8"),
                 InlineKeyboardButton(text="Сен", callback_data="report_month_9")],
                [InlineKeyboardButton(text="Окт", callback_data="report_month_10"),
                 InlineKeyboardButton(text="Ноя", callback_data="report_month_11"),
                 InlineKeyboardButton(text="Дек", callback_data="report_month_12")],
                [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
            ])
            await callback.message.edit_text(f"Выберите месяц для отчета ({year}):", reply_markup=keyboard)
            await callback.answer()
            return
        
        user_state[user_id]["year"] = year
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="⬅️", callback_data="report_year_prev"),
             InlineKeyboardButton(text=str(year), callback_data="report_year_select"),
             InlineKeyboardButton(text="➡️", callback_data="report_year_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
        ])
        await callback.message.edit_text(f"Выберите год для отчета:", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_report_year: {str(e)}")
        await callback.message.edit_text("Произошла ошибка. Попробуйте снова.", reply_markup=None)
 
#Обработчик для отчета (месяц)          
@dp.callback_query(lambda c: c.data.startswith("report_month_"))
async def process_report_month(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        action = callback.data.split("_")[2]
        year = user_state[user_id].get("year", datetime.now().year)

        if action == "prev":
            year -= 1
        elif action == "next":
            year += 1
        elif action in [str(i) for i in range(1, 13)]:
            month = int(action)
            user_state[user_id]["month"] = month
            report_type = user_state[user_id]["report_type"]

            if report_type == "Еженедельный":
                weeks = []
                cal = calendar.Calendar(firstweekday=0)
                for week in cal.monthdayscalendar(year, month):
                    if week[0] != 0:
                        start_day = week[0]
                        end_day = week[6] if week[6] != 0 else monthrange(year, month)[1]
                        if start_day <= monthrange(year, month)[1] and end_day <= monthrange(year, month)[1]:
                            weeks.append((start_day, end_day))

                if not weeks:
                    await callback.message.delete()
                    await callback.message.answer("Нет доступных недель для выбранного месяца.", reply_markup=main_menu)
                    if user_id in user_state:
                        del user_state[user_id]
                    await callback.answer()
                    return

                keyboard = InlineKeyboardMarkup(inline_keyboard=[
                    [InlineKeyboardButton(
                        text=f"Неделя {i+1} ({datetime(year, month, start).strftime('%d.%m')} - {datetime(year, month, end).strftime('%d.%m')})",
                        callback_data=f"report_week_select_{i}"
                    )] for i, (start, end) in enumerate(weeks)
                ] + [[InlineKeyboardButton(text="Назад", callback_data="back_to_report_month")]])

                user_state[user_id]["weeks"] = weeks
                user_state[user_id]["history"].append("report_week_selection")
                await callback.message.edit_text(f"Выберите неделю для отчета ({year}-{month:02d}):", reply_markup=keyboard)
                await callback.answer()
                return

            elif report_type == "Ежемесячный":
                start_date = datetime(year, month, 1).date()
                end_date = datetime(year, month, monthrange(year, month)[1]).date()
                
                start_time = datetime.now()
                logger.info(f"Отправка сообщения 'Идет генерация...' для user_id={user_id}, report_month")
                generating_message = await callback.message.answer("Идет генерация...")
                user_state[user_id]["generating_message_id"] = generating_message.message_id

                await asyncio.sleep(2)

                data = get_monthly_report_data(start_date, end_date)
                doc_buffer = create_monthly_word_report(start_date, end_date, data)

                logger.info(f"Удаление сообщения 'Идет генерация...' для user_id={user_id}, время отображения: {(datetime.now() - start_time).total_seconds()} сек")
                await bot.delete_message(chat_id=callback.message.chat.id, message_id=generating_message.message_id)

                if doc_buffer:
                    filename = f"Ежемесячный_отчет_{start_date.strftime('%Y-%m')}.docx"
                    await callback.message.delete()
                    await bot.send_document(
                        chat_id=callback.message.chat.id,
                        document=BufferedInputFile(doc_buffer.getvalue(), filename=filename),
                        caption=f"Ежемесячный отчет за {calendar.month_name[month]} {year}",
                        reply_markup=main_menu
                    )
                else:
                    await callback.message.delete()
                    await callback.message.answer("Ошибка при создании ежемесячного отчета.", reply_markup=main_menu)

                if user_id in user_state:
                    del user_state[user_id]
                await callback.answer()
                return

        user_state[user_id]["year"] = year
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Янв", callback_data="report_month_1"),
             InlineKeyboardButton(text="Фев", callback_data="report_month_2"),
             InlineKeyboardButton(text="Мар", callback_data="report_month_3")],
            [InlineKeyboardButton(text="Апр", callback_data="report_month_4"),
             InlineKeyboardButton(text="Май", callback_data="report_month_5"),
             InlineKeyboardButton(text="Июн", callback_data="report_month_6")],
            [InlineKeyboardButton(text="Июл", callback_data="report_month_7"),
             InlineKeyboardButton(text="Авг", callback_data="report_month_8"),
             InlineKeyboardButton(text="Сен", callback_data="report_month_9")],
            [InlineKeyboardButton(text="Окт", callback_data="report_month_10"),
             InlineKeyboardButton(text="Ноя", callback_data="report_month_11"),
             InlineKeyboardButton(text="Дек", callback_data="report_month_12")],
            [InlineKeyboardButton(text="⬅️", callback_data="report_month_prev"),
             InlineKeyboardButton(text=str(year), callback_data="report_month_select"),
             InlineKeyboardButton(text="➡️", callback_data="report_month_next")],
            [InlineKeyboardButton(text="Назад", callback_data="back_to_report_menu")]
        ])
        await callback.message.edit_text(f"Выберите месяц для отчета ({year}):", reply_markup=keyboard)
        await callback.answer()
    except Exception as e:
        logger.error(f"Ошибка в process_report_month: {str(e)}")
        if user_id in user_state and "generating_message_id" in user_state[user_id]:
            try:
                await bot.delete_message(chat_id=callback.message.chat.id, 
                                       message_id=user_state[user_id]["generating_message_id"])
            except Exception as delete_error:
                logger.error(f"Ошибка при удалении сообщения 'Идет генерация...': {str(delete_error)}")
        await callback.message.delete()
        await callback.message.answer(f"Произошла ошибка при создании отчета: {str(e)}", reply_markup=main_menu)
        if user_id in user_state:
            del user_state[user_id]
            
#Обработчик для отчета (неделя)          
@dp.callback_query(lambda c: c.data.startswith("report_week_select_"))
async def process_report_week(callback: types.CallbackQuery):
    try:
        user_id = callback.from_user.id
        week_idx = int(callback.data.split("_")[3])
        year = user_state[user_id]["year"]
        month = user_state[user_id]["month"]
        start_day, end_day = user_state[user_id]["weeks"][week_idx]
        start_date = datetime(year, month, start_day).date()
        end_date = datetime(year, month, end_day).date()

        # Сообщение о генерации
        start_time = datetime.now()
        logger.info(f"Отправка сообщения 'Идет генерация...' для user_id={user_id}, report_week")
        generating_message = await callback.message.answer("Идет генерация...")
        user_state[user_id]["generating_message_id"] = generating_message.message_id

        # Получение данных и создание отчета
        data = get_weekly_report_data(start_date, end_date)
        doc_buffer = create_weekly_word_report(start_date, end_date, data)

        # Удаление сообщения "Идет генерация..."
        logger.info(f"Удаление сообщения 'Идет генерация...' для user_id={user_id}, время отображения: {(datetime.now() - start_time).total_seconds()} сек")
        await bot.delete_message(chat_id=callback.message.chat.id, message_id=generating_message.message_id)

        # Отправка документа
        month_name = start_date.strftime("%B")
        doc_name = f"Еженедельный_отчет за_{week_idx+1}_неделю_{month_name}_{year}.docx"
        await callback.message.delete()
        await bot.send_document(
            chat_id=callback.message.chat.id,
            document=BufferedInputFile(doc_buffer.getvalue(), filename=doc_name),
            caption=f"Еженедельный отчет за {week_idx+1}-ю неделю {month_name} {year}",
            reply_markup=main_menu
        )
        logger.info(f"Еженедельный отчет за {week_idx+1}-ю неделю {month_name} {year} создан для user_id={user_id}")

        # Очистка состояния
        if user_id in user_state:
            del user_state[user_id]
        await callback.answer()

    except Exception as e:
        logger.error(f"Ошибка в process_report_week: {str(e)}")
        if user_id in user_state and "generating_message_id" in user_state[user_id]:
            try:
                await bot.delete_message(chat_id=callback.message.chat.id, 
                                       message_id=user_state[user_id]["generating_message_id"])
            except Exception as delete_error:
                logger.error(f"Ошибка при удалении сообщения 'Идет генерация...': {str(delete_error)}")
        await callback.message.delete()
        await callback.message.answer(f"Произошла ошибка при генерации отчета: {str(e)}", reply_markup=main_menu)
        if user_id in user_state:
            del user_state[user_id]

#Обработчик кнопка Помощь            
@dp.message(lambda message: message.text == "Помощь")
async def show_help(message: types.Message):
    try:
        help_text = (
            "🤖 BI Mate — бот-аналитик продаж\n\n"
            "Я автоматически генерирую графики и отчеты по вашим продажам прямо в Telegram.\n\n"
            "📈 Графики\n"
            "— Выберите тип графика (например, 'Динамика выручки').\n"
            "— Укажите период (год, месяц, неделя и т.д.).\n"
            "— Получите график в формате PNG.\n\n"
            "📝 Отчеты\n"
            "— Дашборд: ключевые метрики в формате PDF.\n"
            "— Еженедельный/месячный — Отчёт за неделю/месяц в формате Word.\n\n"
            "📦 Анализ товара\n"
            "— Основные данные и интерактивные графики, выбранного товара\n\n"
            "📊 Анализ продаж\n"
            "— Основные данные и интерактивные графики продаж за выбранный период\n\n"
            "❓ Возникли проблемы? Напишите админу: @foolforu1"
        )
        await message.answer(help_text, reply_markup=main_menu)
    except Exception as e:
        logger.error(f"Ошибка в show_help: {str(e)}")
        await message.answer("Произошла ошибка при отображении помощи. Попробуйте снова.", reply_markup=main_menu)
async def main():
    try:
        logger.info("Запуск бота...")
        await dp.start_polling(bot)
        logger.info("Start polling")
    except Exception as e:
        logger.error(f"Ошибка при запуске бота: {str(e)}")
        await asyncio.sleep(5)
        await main()

if __name__ == '__main__':
    asyncio.run(main())